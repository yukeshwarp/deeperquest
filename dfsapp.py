import streamlit as st
from dotenv import load_dotenv
from writer import report_writer, eval_agent
from planner import plan_research, replanner
from dfs_stepexecutor import execute_step
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup
import markdown as md
import logging
import concurrent.futures

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s"
)

st.set_page_config(
    page_title="Depth-First Deep Research App",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("deepQuest v2")
st.sidebar.title("Research Steps")

def generate_word_doc_from_markdown(markdown_text):
    try:
        html = md.markdown(markdown_text, extensions=['tables'])
        soup = BeautifulSoup(html, "html.parser")
        doc = Document()
        doc.add_heading("DeepQuest Research Report", 0)

        for element in soup.children:
            if element.name and element.name.startswith("h") and element.name[1:].isdigit():
                level = int(element.name[1:])
                doc.add_heading(element.get_text(), level=level)
            elif element.name == "ul":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Bullet")
            elif element.name == "ol":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Number")
            elif element.name == "p":
                doc.add_paragraph(element.get_text())
            elif element.name == "table":
                rows = element.find_all("tr")
                if not rows:
                    continue
                cols = rows[0].find_all(["td", "th"])
                n_cols = len(cols)
                n_rows = len(rows)
                table = doc.add_table(rows=n_rows, cols=n_cols)
                for row_idx, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])
                    for col_idx, cell in enumerate(cells):
                        table.cell(row_idx, col_idx).text = cell.get_text()
                # Add borders to the table
                tbl = table._tbl
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                tblPr = tbl.tblPr
                borders = parse_xml(r'''
                    <w:tblBorders %s>
                        <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                        <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                        <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                    </w:tblBorders>''' % nsdecls('w'))
                tblPr.append(borders)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        logging.error(f"Error converting markdown to Word: {e}")
        return None

# --- Session State Management ---
if "query" not in st.session_state:
    st.session_state.query = ""
if "steps" not in st.session_state:
    st.session_state.steps = []
if "completed_steps" not in st.session_state:
    st.session_state.completed_steps = []
if "context" not in st.session_state:
    st.session_state.context = ""
if "report" not in st.session_state:
    st.session_state.report = None
if "proceed" not in st.session_state:
    st.session_state.proceed = False
if "steps_initialized" not in st.session_state:
    st.session_state.steps_initialized = False

query = st.chat_input("Enter your research query:")
if query and (st.session_state.query != query):
    # Clear all relevant session state for a new query
    st.session_state.steps = []
    st.session_state.completed_steps = []
    st.session_state.context = ""
    st.session_state.report = None
    st.session_state.proceed = False
    st.session_state.steps_initialized = False
    st.session_state.query = query

# Set your max_steps dynamically or statically as needed
max_steps = 20  # Or use a value from Q-learning or user input

# Only generate steps when a new query is submitted
if query and (not st.session_state.steps_initialized or st.session_state.query != query):
    st.session_state.steps = plan_research(query, max_steps=max_steps)
    st.session_state.completed_steps = []
    st.session_state.context = ""
    st.session_state.report = None
    st.session_state.proceed = False
    st.session_state.steps_initialized = True
    st.session_state.query = query
    # st.write(f"Query: {query}")

if st.session_state.query and st.session_state.steps:
    st.write(f"Query: {st.session_state.query}")
    # Always use session state for steps
    if not st.session_state.proceed:
        steps = st.session_state.steps
        completed_steps = st.session_state.completed_steps
        st.sidebar.subheader("Edit Steps Before Proceeding")
        remove_indices = []
        for idx, step in enumerate(steps):
            cols = st.sidebar.columns([8,1])
            clean_step = step.lstrip('.0123456789 ').strip()
            cols[0].write(clean_step)
            if cols[1].button("❌", key=f"remove_{idx}"):
                if len(steps) > 1:
                    steps.pop(idx)
                    st.session_state.steps = steps
                    st.rerun()
                else:
                    st.sidebar.warning("At least one step must remain.")
        # Remove steps marked for deletion
        if remove_indices:
            if len(steps) - len(remove_indices) < 1:
                st.sidebar.warning("At least one step must remain.")
            else:
                for idx in sorted(remove_indices, reverse=True):
                    del steps[idx]
                st.session_state.steps = steps
        # Add new step
        new_step = st.sidebar.text_input("Add a new step", key="add_step")
        if st.sidebar.button("Add Step") and new_step.strip():
            st.session_state.steps.append(new_step.strip())
            st.rerun()
        # Proceed button
        if st.sidebar.button("Proceed with Research"):
            st.session_state.proceed = True
            st.rerun()
        # Show current steps
        # st.sidebar.markdown("**Current Steps:**\n" + "\n".join([step.lstrip('.0123456789 ').strip() for step in st.session_state.steps]))
    else:
        try:
            steps = st.session_state.steps
            completed_steps = st.session_state.completed_steps
            sidebar_steps = st.sidebar.empty()
            # Prepare step display: green tick for completed, plain for pending, no leading dot/number
            completed_step_texts = set(step if isinstance(step, str) else step[0] for step in completed_steps)
            step_lines = []
            for step in steps:
                clean_step = step.lstrip('.0123456789 ').strip()
                if step in completed_step_texts:
                    step_lines.append(f"✅ {clean_step}\n\n")
                else:
                    step_lines.append(f"{clean_step}\n\n")
            sidebar_steps.markdown("\n".join(step_lines))

            context = st.session_state.context
            replan_rounds = 0
            replan_limit_reached = False
            max_steps_warning_shown = False

            progress_bar = st.progress(0, text="Starting research steps...")

            # --- Parallel execution in batches of 3 ---
            batch_size = 3
            i = len(completed_steps)
            while i < len(st.session_state.steps):
                batch_steps = st.session_state.steps[i:i+batch_size]
                with concurrent.futures.ThreadPoolExecutor(max_workers=batch_size) as executor:
                    future_to_step = {executor.submit(execute_step, step, context): step for step in batch_steps}
                    for idx, future in enumerate(concurrent.futures.as_completed(future_to_step)):
                        step = future_to_step[future]
                        try:
                            result = future.result()
                        except Exception as e:
                            logging.error(f"Error executing step '{step}': {e}")
                            st.error("Brain down, try again shortly!")
                            st.stop()
                        completed_steps.append((step, result))
                        context += f"\nStep: {step}\nResult: {result}\n"
                        st.session_state.completed_steps = completed_steps
                        st.session_state.context = context
                        progress = int((len(completed_steps) / len(st.session_state.steps)) * 100)
                        progress_bar.progress(progress / 100, text=f"Completed {len(completed_steps)} of {len(st.session_state.steps)} steps")
                i += batch_size
                # Replanning after every 3 steps
                if not replan_limit_reached:
                    try:
                        new_steps, replan_rounds, replan_limit_reached = replanner(
                            context, st.session_state.steps, replan_rounds, 3, replan_limit_reached, max_steps=max_steps
                        )
                        st.session_state.steps = new_steps
                    except Exception as e:
                        logging.error(f"Error during replanning: {e}")
                        st.error("Brain down, try again shortly!")
                        st.stop()

            progress_bar.progress(1.0, text="All steps completed!")

            # Generate report only if not already in session state
            if not st.session_state.report:
                try:
                    st.session_state.report = report_writer(context)
                except Exception as e:
                    logging.error(f"Error generating report: {e}")
                    st.error("Brain down, try again shortly!")
                    st.stop()

        except Exception as e:
            logging.critical(f"Critical error in main UI: {e}")
            st.error("Brain down, try again shortly!")

# --- Always display report and download button if available ---
if st.session_state.report:
    # st.write(f"Query: {query}")
    st.subheader("Final Research Report")
    st.markdown(st.session_state.report)
    word_buffer = generate_word_doc_from_markdown(st.session_state.report)
    if word_buffer:
        st.download_button(
            label="Download Report as Word Document",
            data=word_buffer,
            file_name="deepquest_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        st.error("Brain down, try again shortly!")
